import argparse
import os
import requests
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from rich.console import Console
from rich.panel import Panel

# Inicializa o console do Rich para exibir mensagens formatadas
console = Console()

# Função para converter texto em áudio usando a API OpenAI TTS (Create Speech)
def text_to_speech_openai(text, output_audio_path, model="tts-1", voice="alloy", response_format="mp3"):
    url = "https://api.openai.com/v1/audio/speech"
    
    # Configura os cabeçalhos de autenticação
    headers = {
        "Authorization": f"Bearer {os.getenv('OPENAI_API_KEY')}",
        "Content-Type": "application/json"
    }

    # Define o corpo da requisição com o modelo TTS, voz e formato de saída
    data = {
        "model": model,
        "input": text,
        "voice": voice,
        "response_format": response_format
    }

    # Envia a solicitação para a API da OpenAI
    response = requests.post(url, headers=headers, json=data)
    
    # Verifica se a solicitação foi bem-sucedida
    if response.status_code == 200:
        # Salva o conteúdo de áudio gerado no arquivo de saída
        with open(output_audio_path, "wb") as audio_file:
            audio_file.write(response.content)
        console.print(f"[green]Áudio gerado e salvo em: {output_audio_path}[/green]")
    else:
        # Exibe a mensagem de erro em caso de falha
        console.print(f"[red]Erro ao gerar o áudio: {response.status_code} - {response.text}[/red]")
        return False
    return True

# Função para enviar áudio para a OpenAI e obter a transcrição
def transcribe_audio(file_path):
    client = OpenAI()

    # Abre o arquivo de áudio no modo de leitura binária
    with open(file_path, "rb") as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    return transcription.text

# Função para enviar áudio para a OpenAI e obter a tradução para o inglês
def translate_audio(file_path):
    client = OpenAI()

    # Abre o arquivo de áudio no modo de leitura binária
    with open(file_path, "rb") as audio_file:
        translation = client.audio.translations.create(
            model="whisper-1",
            file=audio_file
        )
    return translation.text

# Função para salvar transcrição e tradução em um arquivo Word (.docx)
def export_to_docx(transcription_text, translation_text, output_path="transcription_and_translation.docx"):
    doc = Document()

    doc.add_heading('Transcrição no idioma original:', level=1)
    doc.add_paragraph(transcription_text)

    doc.add_heading('Tradução para o inglês:', level=1)
    doc.add_paragraph(translation_text)

    doc.save(output_path)
    console.print(f"[green]Arquivo salvo com sucesso em: {output_path}[/green]")

if __name__ == "__main__":
    load_dotenv()  # Carrega as variáveis de ambiente do arquivo .env (como a chave da API da OpenAI)

    # Configuração dos argumentos da linha de comando com opções curtas
    parser = argparse.ArgumentParser(description="Transformar texto em áudio, transcrever e traduzir")
    
    # Opção curta -f para o arquivo de texto de entrada
    parser.add_argument("-f", "--file", type=str, required=True, help="Caminho para o arquivo de texto de entrada")
    
    # Opção curta -o para o arquivo de áudio gerado
    parser.add_argument("-o", "--output_audio", type=str, default="output_audio.mp3", help="Caminho para o arquivo de áudio gerado")
    
    # Opção curta -d para o arquivo .docx gerado
    parser.add_argument("-d", "--output_docx", type=str, default="transcription_and_translation.docx", help="Nome do arquivo de saída .docx")

    # Adiciona opções para o modelo de fala e a voz
    parser.add_argument("--model", type=str, default="tts-1", help="Modelo de TTS da OpenAI (padrão: tts-1)")
    parser.add_argument("--voice", type=str, default="alloy", help="Voz a ser usada para gerar o áudio (padrão: alloy)")
    
    args = parser.parse_args()

    # Passo 1: Ler o arquivo de texto de entrada
    with open(args.file, "r", encoding="utf-8") as f:
        input_text = f.read()

    # Passo 2: Converter texto em áudio usando a OpenAI API
    audio_generated = text_to_speech_openai(input_text, args.output_audio, model=args.model, voice=args.voice)

    if audio_generated:
        # Passo 3: Transcrever o áudio gerado
        transcription_text = transcribe_audio(args.output_audio)
        console.print(Panel.fit(f"[bold]Transcrição no idioma original:[/bold]\n{transcription_text}",
                                title="Transcrição", subtitle="Idioma Original", border_style="blue"))

        # Passo 4: Traduzir o áudio para inglês
        translation_text = translate_audio(args.output_audio)
        console.print(Panel.fit(f"[bold]Tradução para o inglês:[/bold]\n{translation_text}",
                                title="Tradução", subtitle="EN-US", border_style="green"))

        # Passo 5: Exportar a transcrição e tradução para um arquivo .docx
        export_to_docx(transcription_text, translation_text, args.output_docx)
